"""Report exporter for FlowTrack.

Generates Word (.docx) documents from weekly summary data using python-docx.
"""

import logging
import os
from datetime import timedelta

from flowtrack.core.models import CategorySummary, DailySummary, WeeklySummary
from flowtrack.reporting.formatter import TextFormatter

logger = logging.getLogger(__name__)


class ReportExporter:
    """Exports weekly summary data to a formatted Word document (.docx)."""

    def export_weekly(
        self, summary: WeeklySummary, user_name: str, output_path: str
    ) -> str:
        """Generate a .docx file from weekly summary data.

        Args:
            summary: The weekly summary to export.
            user_name: The user's configured display name.
            output_path: File path for the generated .docx file.

        Returns:
            The path to the generated file.

        Raises:
            ImportError: If python-docx is not installed.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for report export. "
                "Install it with: pip install python-docx"
            )

        # Create parent directories if needed
        parent_dir = os.path.dirname(output_path)
        if parent_dir:
            os.makedirs(parent_dir, exist_ok=True)

        doc = Document()

        # --- Title page ---
        self._add_title_page(doc, summary, user_name)

        # --- Weekly category summary table ---
        doc.add_heading("Weekly Summary", level=1)
        self._add_category_table(doc, summary.categories, summary.total_time, summary.total_sessions)

        # --- Day-by-day breakdown ---
        doc.add_heading("Daily Breakdown", level=1)
        for daily in summary.daily_breakdowns:
            day_label = daily.date.strftime("%A, %B %d, %Y")
            doc.add_heading(day_label, level=2)
            if not daily.categories:
                doc.add_paragraph("No activity recorded.")
            else:
                self._add_category_table(
                    doc, daily.categories, daily.total_time, daily.total_sessions
                )

        doc.save(output_path)
        return output_path

    def _add_title_page(
        self, doc, summary: WeeklySummary, user_name: str
    ) -> None:
        """Add a title page with report title, date range, and user name."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("FlowTrack Weekly Report")
        run.bold = True
        run.font.size = Pt(24)

        # Date range
        start_str = summary.start_date.strftime("%B %d, %Y")
        end_str = summary.end_date.strftime("%B %d, %Y")
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"{start_str} - {end_str}")
        run.font.size = Pt(14)

        # User name
        if user_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = name_para.add_run(f"Prepared for: {user_name}")
            run.font.size = Pt(12)

        # Page break after title page
        doc.add_page_break()

    def _add_category_table(
        self,
        doc,
        categories: list[CategorySummary],
        total_time: timedelta,
        total_sessions: int,
    ) -> None:
        """Add a category summary table to the document."""
        from docx.shared import Pt

        # Header row + data rows + total row
        table = doc.add_table(rows=1 + len(categories) + 1, cols=3)
        table.style = "Light Grid Accent 1"

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Category"
        header_cells[1].text = "Total Time"
        header_cells[2].text = "Sessions"

        # Data rows
        for i, cat in enumerate(categories, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = cat.category
            row_cells[1].text = TextFormatter.format_duration(cat.total_time)
            row_cells[2].text = str(cat.completed_sessions)

        # Total row
        total_row = table.rows[-1].cells
        total_row[0].text = "Total"
        total_row[1].text = TextFormatter.format_duration(total_time)
        total_row[2].text = str(total_sessions)

        # Bold the header and total rows
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for cell in table.rows[-1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()  # spacing after table
