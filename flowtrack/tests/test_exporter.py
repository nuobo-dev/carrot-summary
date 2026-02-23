"""Tests for the ReportExporter class."""

import os
import tempfile
from datetime import date, timedelta

import pytest
from docx import Document

from flowtrack.core.models import CategorySummary, DailySummary, WeeklySummary
from flowtrack.reporting.exporter import ReportExporter


def _make_category(name: str, minutes: int, sessions: int) -> CategorySummary:
    return CategorySummary(
        category=name,
        sub_categories={},
        total_time=timedelta(minutes=minutes),
        completed_sessions=sessions,
    )


def _make_weekly_summary() -> WeeklySummary:
    """Build a realistic weekly summary for testing."""
    start = date(2025, 1, 6)  # Monday
    end = date(2025, 1, 12)   # Sunday

    dev_cat = _make_category("Development", 480, 8)
    email_cat = _make_category("Email & Communication", 120, 3)
    weekly_cats = [dev_cat, email_cat]

    daily_breakdowns = []
    for i in range(7):
        d = date(2025, 1, 6 + i)
        if i < 5:  # weekdays have activity
            daily_breakdowns.append(
                DailySummary(
                    date=d,
                    categories=[
                        _make_category("Development", 96, 2),
                        _make_category("Email & Communication", 24, 1),
                    ],
                    total_time=timedelta(minutes=120),
                    total_sessions=3,
                )
            )
        else:  # weekends empty
            daily_breakdowns.append(
                DailySummary(date=d, categories=[], total_time=timedelta(), total_sessions=0)
            )

    return WeeklySummary(
        start_date=start,
        end_date=end,
        daily_breakdowns=daily_breakdowns,
        categories=weekly_cats,
        total_time=timedelta(minutes=600),
        total_sessions=11,
    )


class TestReportExporter:
    """Tests for ReportExporter.export_weekly."""

    def test_creates_docx_file(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        result = exporter.export_weekly(summary, "Alice", out)

        assert result == out
        assert os.path.isfile(out)

    def test_creates_parent_directories(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "nested" / "dir" / "report.docx")

        exporter.export_weekly(summary, "Bob", out)

        assert os.path.isfile(out)

    def test_title_page_content(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "Alice Smith", out)

        doc = Document(out)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "FlowTrack Weekly Report" in all_text
        assert "January 06, 2025" in all_text
        assert "January 12, 2025" in all_text
        assert "Alice Smith" in all_text

    def test_title_page_no_user_name(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "", out)

        doc = Document(out)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "FlowTrack Weekly Report" in all_text
        assert "Prepared for:" not in all_text

    def test_weekly_summary_table(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "Alice", out)

        doc = Document(out)
        # First table is the weekly summary table
        assert len(doc.tables) >= 1
        table = doc.tables[0]

        # Header row
        assert table.rows[0].cells[0].text == "Category"
        assert table.rows[0].cells[1].text == "Total Time"
        assert table.rows[0].cells[2].text == "Sessions"

        # Data rows
        assert table.rows[1].cells[0].text == "Development"
        assert table.rows[1].cells[1].text == "8h 0m"
        assert table.rows[1].cells[2].text == "8"

        assert table.rows[2].cells[0].text == "Email & Communication"
        assert table.rows[2].cells[1].text == "2h 0m"
        assert table.rows[2].cells[2].text == "3"

        # Total row
        last_row = table.rows[-1]
        assert last_row.cells[0].text == "Total"
        assert last_row.cells[1].text == "10h 0m"
        assert last_row.cells[2].text == "11"

    def test_daily_breakdown_sections(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "Alice", out)

        doc = Document(out)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Check day headings appear
        assert "Monday, January 06, 2025" in all_text
        assert "Sunday, January 12, 2025" in all_text

        # Weekend days should show "No activity recorded."
        assert "No activity recorded." in all_text

    def test_daily_tables_present(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "Alice", out)

        doc = Document(out)
        # 1 weekly table + 5 weekday tables = 6 tables total
        # (weekends have no tables, just "No activity" text)
        assert len(doc.tables) == 6

    def test_headings_present(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "report.docx")

        exporter.export_weekly(summary, "Alice", out)

        doc = Document(out)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

        assert "Weekly Summary" in headings
        assert "Daily Breakdown" in headings

    def test_empty_weekly_summary(self, tmp_path):
        exporter = ReportExporter()
        summary = WeeklySummary(
            start_date=date(2025, 1, 6),
            end_date=date(2025, 1, 12),
            daily_breakdowns=[
                DailySummary(date=date(2025, 1, 6 + i))
                for i in range(7)
            ],
            categories=[],
            total_time=timedelta(),
            total_sessions=0,
        )
        out = str(tmp_path / "report.docx")

        result = exporter.export_weekly(summary, "Alice", out)

        assert os.path.isfile(result)
        doc = Document(out)
        # Weekly table with just header + total, no data rows
        assert len(doc.tables) >= 1

    def test_returns_output_path(self, tmp_path):
        exporter = ReportExporter()
        summary = _make_weekly_summary()
        out = str(tmp_path / "my_report.docx")

        result = exporter.export_weekly(summary, "Test", out)

        assert result == out
